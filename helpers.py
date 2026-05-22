# helpers.py — вспомогательные функции для генерации ВКР
from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def init_document():
    """Создание документа с базовыми настройками стилей."""
    doc = Document()

    # Настройка стиля Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Принудительно задаём шрифт через XML (rFonts)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" '
                           f'w:hAnsi="Times New Roman" w:cs="Times New Roman" '
                           f'w:eastAsia="Times New Roman"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:cs'), 'Times New Roman')
        rfonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Поля первой секции
    section = doc.sections[0]
    set_page_margins(section)

    return doc


def set_page_margins(section):
    """Установка полей страницы: лево 30мм, право 15мм, верх/низ 20мм."""
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def set_run_font(run, size=14, bold=False, italic=False, font_name='Times New Roman'):
    """Форматирование run с принудительным rFonts через XML."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name

    # Принудительно задаём шрифт через XML
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" '
                           f'w:hAnsi="{font_name}" w:cs="{font_name}" '
                           f'w:eastAsia="{font_name}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
        rfonts.set(qn('w:cs'), font_name)
        rfonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(1.25),
                         space_before=Pt(0), space_after=Pt(0),
                         line_spacing=1.5, keep_together=False,
                         page_break_before=False):
    """Форматирование абзаца."""
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = first_line_indent
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.keep_together = keep_together
    pf.page_break_before = page_break_before


def add_body_text(doc, text):
    """Добавление абзаца основного текста."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14)
    set_paragraph_format(p)
    return p


def add_heading_section(doc, text):
    """Заголовок раздела: ПРОПИСНЫЕ, жирный, по центру, с новой страницы."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_run_font(run, size=14, bold=True)
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0),
        page_break_before=True,
        space_after=Pt(12)
    )
    return p


def add_heading_sub(doc, text):
    """Заголовок подраздела: жирный, по центру, без новой страницы, отступ сверху 12 пт."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0),
        space_before=Pt(12),
        space_after=Pt(6)
    )
    return p


def add_table_caption(doc, number, title):
    """Подпись таблицы: Таблица N — Название, по центру, 14 пт."""
    p = doc.add_paragraph()
    run = p.add_run(f'Таблица {number} \u2014 {title}')
    set_run_font(run, size=14)
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0),
        space_before=Pt(12),
        space_after=Pt(6)
    )
    return p


def make_table(doc, headers, rows):
    """Создание таблицы со стилем Table Grid."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=12, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Отступ после таблицы
    p_after = doc.add_paragraph()
    set_paragraph_format(p_after, space_before=Pt(6))

    return table


def add_figure(doc, number, caption, image_path=None):
    """Вставка рисунка-заглушки с подписью."""
    # Абзац-заглушка вместо реального изображения
    p_img = doc.add_paragraph()
    if image_path:
        try:
            run = p_img.add_run()
            run.add_picture(image_path, width=Cm(15))
        except Exception:
            run = p_img.add_run(f'[Здесь размещается Рисунок {number}]')
            set_run_font(run, size=12, italic=True)
    else:
        run = p_img.add_run(f'[Здесь размещается Рисунок {number}]')
        set_run_font(run, size=12, italic=True)
    set_paragraph_format(
        p_img,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0),
        space_before=Pt(12)
    )

    # Подпись под рисунком
    p_cap = doc.add_paragraph()
    run_cap = p_cap.add_run(f'Рисунок {number} \u2014 {caption}')
    set_run_font(run_cap, size=12)
    set_paragraph_format(
        p_cap,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0),
        space_after=Pt(6)
    )
    return p_img


def add_dash_item(doc, text, is_last=False):
    """Пункт списка с тире. Последний заканчивается точкой, остальные — точкой с запятой."""
    ending = '.' if is_last else ';'
    p = doc.add_paragraph()
    run = p.add_run(f'\u2013 {text}{ending}')
    set_run_font(run, size=14)
    set_paragraph_format(p)
    return p


def add_page_numbers(doc):
    """Нумерация страниц: пропуск 1-й секции (содержание), нумерация со 2-й секции с номера 3."""
    # Убираем нумерацию из первой секции
    section0 = doc.sections[0]
    footer0 = section0.footer
    footer0.is_linked_to_previous = False
    # Очистка footer первой секции
    for p in footer0.paragraphs:
        p.text = ''

    # Нумерация во второй секции
    if len(doc.sections) > 1:
        section1 = doc.sections[1]

        # Установка начала нумерации с 3
        sectPr = section1._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="3"/>')
            sectPr.append(pgNumType)
        else:
            pgNumType.set(qn('w:start'), '3')

        footer1 = section1.footer
        footer1.is_linked_to_previous = False

        # Очищаем существующие параграфы
        for p in footer1.paragraphs:
            p._element.getparent().remove(p._element)

        # Создаём новый параграф с номером страницы
        p = footer1.paragraphs[0] if footer1.paragraphs else footer1.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        set_run_font(run, size=9)

        # Вставляем поле PAGE
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._element.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._element.append(fldChar2)


def add_toc_line(doc, text, page_num):
    """Строка ручного оглавления с точечным заполнителем и номером страницы."""
    p = doc.add_paragraph()

    # Определяем отступ по уровню
    is_sub = any(text.startswith(f'{i}.') for i in range(1, 10))
    indent = Cm(1.25) if is_sub else Cm(0)

    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5
    )

    # Добавляем табуляцию с точечным заполнителем
    pPr = p._element.get_or_add_pPr()
    tabs = parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        f'  <w:tab w:val="right" w:leader="dot" w:pos="9356"/>'
        f'</w:tabs>'
    )
    pPr.append(tabs)

    run = p.add_run(text)
    set_run_font(run, size=14)

    run_tab = p.add_run('\t')
    set_run_font(run_tab, size=14)

    run_num = p.add_run(str(page_num))
    set_run_font(run_num, size=14)

    return p
